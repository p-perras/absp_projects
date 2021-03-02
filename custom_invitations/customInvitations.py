#! python3
# customInvitations.py
# ABSP - Chapter 15

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def create_invitations(guestList, docName):
    """
    Summary:
        Creates invitations based on names in a given text file.

    Args:
        guestList (str): Text file containing names.
        docName (str): Doc file to save invitations in.
    """

    doc = docx.Document()

    greeting = 'It would be a pleasure to have the company of'
    address = 'at 11101 Memory lane on the evening of'
    date = 'April 31st'
    time = "at 7 O'Clock"

    namesFile = open(guestList, 'r')
    for guest in namesFile.readlines():
        name = guest[:-1]

        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(greeting)
        r1.font.bold = True
        r1.font.italic = True
        r1.font.size = Pt(13)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(name)
        r2.font.bold = True
        r2.font.size = Pt(15)

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(address)
        r3.font.bold = True
        r3.font.italic = True
        r3.font.size = Pt(12)

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(date)
        r4.font.size = Pt(12)

        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r5 = p5.add_run(time)
        r5.font.bold = True
        r5.font.italic = True
        r5.font.size = Pt(12)

        doc.add_page_break()

    doc.save(docName)


if __name__ == '__main__':
    create_invitations('guests.txt', 'invitations.docx')

